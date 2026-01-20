import streamlit as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import json

class PortfolioBuilder:
    def __init__(self):
        self.templates = {
            "Designer Portfolio": self.build_designer_portfolio,
            "Dark Modern Portfolio": self.build_dark_modern_portfolio,
            "Neon Cyberpunk": self.build_neon_cyberpunk,
            "Elegant Minimalist": self.build_elegant_minimalist,
        }
    
    def _format_list_items(self, items):
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []
    
    def generate_portfolio(self, data, template_name):
        doc = Document()
        if template_name in self.templates:
            doc = self.templates[template_name](doc, data)
        else:
            doc = self.build_modern_developer(doc, data)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def build_modern_developer(self, doc, data):
        styles = doc.styles
        name_style = styles.add_style('MD Name', 1) if 'MD Name' not in styles else styles['MD Name']
        name_style.font.size = Pt(32)
        name_style.font.bold = True
        name_style.font.color.rgb = RGBColor(0, 122, 204)
        name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(data['personal_info']['name'].upper(), style=name_style)
        doc.add_paragraph(data['personal_info']['tagline']).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_parts = [
            data['personal_info'].get('email', ''),
            data['personal_info'].get('phone', ''),
            data['personal_info'].get('linkedin', ''),
            data['personal_info'].get('github', '')
        ]
        contact.add_run(' | '.join([p for p in contact_parts if p]))
        
        if data.get('about'):
            doc.add_paragraph('ABOUT ME').runs[0].font.bold = True
            doc.add_paragraph(data['about'])
        
        if data.get('skills'):
            doc.add_paragraph('SKILLS').runs[0].font.bold = True
            for category, skills_list in data['skills'].items():
                p = doc.add_paragraph()
                p.add_run(f"{category}: ").bold = True
                p.add_run(' • '.join(self._format_list_items(skills_list)))
        
        if data.get('projects'):
            doc.add_paragraph('PROJECTS').runs[0].font.bold = True
            for proj in data['projects']:
                p = doc.add_paragraph()
                p.add_run(proj['title']).bold = True
                if proj.get('tech_stack'):
                    p.add_run(f" | {proj['tech_stack']}")
                doc.add_paragraph(proj['description']).paragraph_format.left_indent = Inches(0.3)
        
        if data.get('experience'):
            doc.add_paragraph('EXPERIENCE').runs[0].font.bold = True
            for exp in data['experience']:
                p = doc.add_paragraph()
                p.add_run(exp['role']).bold = True
                if exp.get('company'):
                    p.add_run(f" | {exp['company']}")
                if exp.get('duration'):
                    p.add_run(f" | {exp['duration']}")
                if exp.get('description'):
                    doc.add_paragraph(exp['description']).paragraph_format.left_indent = Inches(0.3)
        
        if data.get('education'):
            doc.add_paragraph('EDUCATION').runs[0].font.bold = True
            for edu in data['education']:
                p = doc.add_paragraph()
                p.add_run(edu['degree']).bold = True
                details = []
                if edu.get('college'):
                    details.append(edu['college'])
                if edu.get('cgpa'):
                    details.append(f"CGPA: {edu['cgpa']}")
                if edu.get('duration'):
                    details.append(edu['duration'])
                if details:
                    doc.add_paragraph(' | '.join(details)).paragraph_format.left_indent = Inches(0.3)
        
        if data.get('certifications'):
            doc.add_paragraph('CERTIFICATIONS').runs[0].font.bold = True
            for cert in data['certifications']:
                p = doc.add_paragraph()
                p.add_run(f"• {cert['name']}")
                if cert.get('issuer'):
                    p.add_run(f" - {cert['issuer']}")
        
        if data.get('achievements'):
            doc.add_paragraph('ACHIEVEMENTS').runs[0].font.bold = True
            for achievement in data['achievements']:
                doc.add_paragraph(f"• {achievement}")
        
        doc.add_paragraph('CONTACT').runs[0].font.bold = True
        contact_info = []
        if data['personal_info'].get('email'):
            contact_info.append(f"Email: {data['personal_info']['email']}")
        if data['personal_info'].get('linkedin'):
            contact_info.append(f"LinkedIn: {data['personal_info']['linkedin']}")
        if data['personal_info'].get('github'):
            contact_info.append(f"GitHub: {data['personal_info']['github']}")
        if data['personal_info'].get('twitter'):
            contact_info.append(f"Twitter: {data['personal_info']['twitter']}")
        for info in contact_info:
            doc.add_paragraph(f"• {info}")
        
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        return doc
  
    def build_designer_portfolio(self, doc, data):
        return self.build_modern_developer(doc, data)
    
    def build_dark_modern_portfolio(self, doc, data):
        return self.build_modern_developer(doc, data)
    
    def build_neon_cyberpunk(self, doc, data):
        return self.build_modern_developer(doc, data)
    
    def build_elegant_minimalist(self, doc, data):
        return self.build_modern_developer(doc, data)

def generate_preview_html(data, template_name):
    # Format projects
    projects_html = ""
    for i, proj in enumerate(data.get('projects', []), 1):
        if template_name == "Designer Portfolio":
            projects_html += f"""
            <div class="project-card">
                <div class="project-header">
                    <h3 class="project-title">{proj['title']}</h3>
                </div>
                <div class="project-body">
                    <div style="color: #ff6b6b; font-weight: 600; margin-bottom: 15px;">{proj.get('tech_stack', '')}</div>
                    <p style="color: #555; line-height: 1.8; margin-bottom: 20px;">{proj['description']}</p>
                    {f'<a href="{proj["link"]}" style="color: #ff6b6b; font-weight: 600; text-decoration: none;">View Project →</a>' if proj.get('link') else ''}
                </div>
            </div>
            """
        else:
            projects_html += f"""
            <div class="project-card">
                <div class="project-num">PROJECT {i:02d}</div>
                <h3 class="project-title">{proj['title']}</h3>
                <div class="project-tech">{proj.get('tech_stack', '')}</div>
                <p class="project-desc">{proj['description']}</p>
                {f'<a href="{proj["link"]}" class="project-link">View Project →</a>' if proj.get('link') else ''}
            </div>
            """
    
    # Format experience
    experience_html = ""
    for exp in data.get('experience', []):
        experience_html += f"""
        <div class="exp-item">
            <div class="exp-role">{exp['role']}</div>
            <div class="exp-company">{exp.get('company', '')}</div>
            <div class="exp-duration">{exp.get('duration', '')}</div>
            <p class="exp-desc">{exp.get('description', '')}</p>
        </div>
        """
    
    # Format education
    education_html = ""
    for edu in data.get('education', []):
        education_html += f"""
        <div class="edu-card">
            <div class="edu-degree">{edu['degree']}</div>
            <div class="edu-college">{edu.get('college', '')}</div>
            <div class="edu-details">
                <span>{edu.get('cgpa', '')}</span>
                <span>{edu.get('duration', '')}</span>
            </div>
        </div>
        """
    
    # Format certifications
    certifications_html = ""
    for cert in data.get('certifications', []):
        certifications_html += f"""
        <div class="cert-item">
            <div class="cert-name">{cert['name']}</div>
            <div class="cert-issuer">{cert.get('issuer', '')}</div>
        </div>
        """
    
    # Format achievements
    achievements_html = ""
    for achievement in data.get('achievements', []):
        achievements_html += f'<div class="achievement-item">• {achievement}</div>'
    
    templates = {
        "Designer Portfolio": f"""
        <!DOCTYPE html>
        <html><head><meta charset="UTF-8">
        <link href="https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;600;700;900&display=swap" rel="stylesheet">
        <style>
            * {{ margin: 0; padding: 0; box-sizing: border-box; }}
            body {{ font-family: 'Poppins', sans-serif; background: #ffffff; }}
            nav {{ background: white; padding: 25px 60px; display: flex; justify-content: space-between; align-items: center; box-shadow: 0 4px 20px rgba(0,0,0,0.08); position: sticky; top: 0; z-index: 1000; }}
            .logo {{ font-size: 1.8em; font-weight: 900; background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }}
            .nav-links {{ display: flex; gap: 40px; list-style: none; align-items: center; }}
            .nav-links a {{ color: #333; text-decoration: none; font-weight: 600; transition: color 0.3s; }}
            .nav-links a:hover {{ color: #ff6b6b; }}
            .cta-btn {{ background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); color: white; padding: 12px 30px; border-radius: 50px; text-decoration: none; font-weight: 600; box-shadow: 0 5px 20px rgba(255, 107, 107, 0.3); transition: all 0.3s; }}
            .cta-btn:hover {{ transform: translateY(-3px); box-shadow: 0 8px 30px rgba(255, 107, 107, 0.4); }}
            .hero {{ background: linear-gradient(135deg, #f8f9fa 0%, #ffffff 100%); padding: 120px 80px; text-align: center; }}
            .hero-tag {{ color: #ff6b6b; font-size: 1em; text-transform: uppercase; letter-spacing: 3px; font-weight: 700; margin-bottom: 25px; }}
            .hero h1 {{ font-size: 4.5em; color: #1a1a2e; margin-bottom: 30px; line-height: 1.1; font-weight: 900; }}
            .hero-desc {{ color: #666; line-height: 2; font-size: 1.2em; margin-bottom: 30px; max-width: 800px; margin-left: auto; margin-right: auto; }}
            .hero-btn {{ display: inline-block; background: #1a1a2e; color: white; padding: 15px 40px; border-radius: 8px; text-decoration: none; font-weight: 600; transition: all 0.3s; margin-top: 20px; }}
            .hero-btn:hover {{ background: #333; transform: translateX(5px); }}
            .section {{ padding: 100px 80px; }}
            .section-alt {{ background: #f8f9fa; }}
            .section-title {{ text-align: center; font-size: 3.5em; color: #1a1a2e; margin-bottom: 70px; font-weight: 900; }}
            .section-title::after {{ content: ''; display: block; width: 100px; height: 5px; background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); margin: 20px auto 0; border-radius: 3px; }}
            .skills-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(250px, 1fr)); gap: 30px; max-width: 1200px; margin: 0 auto; }}
            .skill-card {{ background: white; padding: 30px; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.08); }}
            .skill-category {{ font-size: 1.3em; font-weight: 700; color: #ff6b6b; margin-bottom: 15px; }}
            .skill-list {{ color: #666; line-height: 2; }}
            .projects-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(400px, 1fr)); gap: 40px; max-width: 1400px; margin: 0 auto; }}
            .project-card {{ background: white; border-radius: 20px; overflow: hidden; box-shadow: 0 10px 40px rgba(0,0,0,0.1); transition: all 0.3s; }}
            .project-card:hover {{ transform: translateY(-10px); box-shadow: 0 20px 60px rgba(0,0,0,0.15); }}
            .project-header {{ background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); padding: 30px; color: white; }}
            .project-title {{ font-size: 1.8em; font-weight: 700; margin-bottom: 10px; }}
            .project-body {{ padding: 30px; }}
            .exp-item {{ background: white; padding: 35px; border-radius: 15px; box-shadow: 0 5px 20px rgba(0,0,0,0.08); margin-bottom: 30px; border-left: 5px solid #ff6b6b; max-width: 900px; margin-left: auto; margin-right: auto; }}
            .exp-role {{ font-size: 1.6em; font-weight: 700; color: #1a1a2e; margin-bottom: 10px; }}
            .exp-company {{ color: #ff6b6b; font-weight: 600; margin-bottom: 5px; }}
            .exp-duration {{ color: #999; font-size: 0.9em; margin-bottom: 15px; }}
            .exp-desc {{ color: #666; line-height: 1.8; }}
            .edu-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(350px, 1fr)); gap: 30px; max-width: 1200px; margin: 0 auto; }}
            .edu-card {{ background: white; padding: 35px; border-radius: 15px; box-shadow: 0 5px 20px rgba(0,0,0,0.08); border-top: 4px solid #feca57; }}
            .edu-degree {{ font-size: 1.4em; font-weight: 700; color: #1a1a2e; margin-bottom: 10px; }}
            .edu-college {{ color: #666; margin-bottom: 10px; }}
            .edu-details {{ display: flex; gap: 20px; color: #999; font-size: 0.95em; }}
            .cert-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(300px, 1fr)); gap: 25px; max-width: 1200px; margin: 0 auto; }}
            .cert-item {{ background: white; padding: 25px; border-radius: 12px; box-shadow: 0 5px 15px rgba(0,0,0,0.08); border-left: 4px solid #ff6b6b; }}
            .cert-name {{ font-weight: 700; color: #1a1a2e; margin-bottom: 8px; }}
            .cert-issuer {{ color: #ff6b6b; font-size: 0.9em; }}
            .achievement-item {{ background: white; padding: 25px 30px; border-radius: 12px; box-shadow: 0 5px 15px rgba(0,0,0,0.08); margin-bottom: 20px; color: #555; font-size: 1.05em; line-height: 1.8; max-width: 900px; margin-left: auto; margin-right: auto; }}
            .contact-section {{ background: linear-gradient(135deg, #1a1a2e 0%, #2d3561 100%); color: white; text-align: center; }}
            .contact-grid {{ display: flex; justify-content: center; gap: 40px; margin-top: 50px; flex-wrap: wrap; }}
            .contact-item {{ background: rgba(255,255,255,0.1); padding: 25px 35px; border-radius: 12px; backdrop-filter: blur(10px); transition: all 0.3s; }}
            .contact-item:hover {{ background: rgba(255,255,255,0.2); transform: translateY(-5px); }}
            .contact-label {{ font-size: 0.9em; opacity: 0.8; margin-bottom: 8px; }}
            .contact-value {{ font-weight: 600; font-size: 1.1em; word-break: break-word; }}
            .download-btn {{ display: inline-block; background: linear-gradient(135deg, #ff6b6b 0%, #feca57 100%); color: white; padding: 18px 45px; border-radius: 50px; text-decoration: none; font-weight: 700; font-size: 1.1em; margin-top: 40px; box-shadow: 0 10px 30px rgba(255, 107, 107, 0.3); transition: all 0.3s; }}
            .download-btn:hover {{ transform: translateY(-3px); box-shadow: 0 15px 40px rgba(255, 107, 107, 0.4); }}
        </style></head><body>
            <nav>
                <div class="logo">PORTFOLIO</div>
                <ul class="nav-links">
                    <li><a href="#about">About</a></li>
                    <li><a href="#skills">Skills</a></li>
                    <li><a href="#projects">Projects</a></li>
                    <li><a href="#experience">Experience</a></li>
                    <li><a href="#education">Education</a></li>
                    <li><a href="#contact">Contact</a></li>
                </ul>
                <a href="mailto:{data['personal_info']['email']}" class="cta-btn">Get In Touch</a>
            </nav>
            
            <section class="hero" id="about">
                <div class="hero-tag">WELCOME TO MY PORTFOLIO</div>
                <h1>{data['personal_info']['name']}</h1>
                <p style="font-size: 1.5em; color: #ff6b6b; margin-bottom: 30px;">{data['personal_info']['tagline']}</p>
                <p class="hero-desc">{data['about']}</p>
                <a href="#contact" class="hero-btn">Let's Connect →</a>
            </section>
            
            <section class="section section-alt" id="skills">
                <h2 class="section-title">Skills & Expertise</h2>
                <div class="skills-grid">
                    {''.join([f'<div class="skill-card"><div class="skill-category">{cat}</div><div class="skill-list">{", ".join(skills)}</div></div>' for cat, skills in data.get('skills', {}).items()])}
                </div>
            </section>
            
            <section class="section" id="projects">
                <h2 class="section-title">Featured Projects</h2>
                <div class="projects-grid">{projects_html}</div>
            </section>
            
            <section class="section section-alt" id="experience">
                <h2 class="section-title">Experience</h2>
                {experience_html}
            </section>
            
            <section class="section" id="education">
                <h2 class="section-title">Education</h2>
                <div class="edu-grid">{education_html}</div>
            </section>
            
            <section class="section section-alt" id="certifications">
                <h2 class="section-title">Certifications</h2>
                <div class="cert-grid">{certifications_html}</div>
            </section>
            
            <section class="section" id="achievements">
                <h2 class="section-title">Achievements</h2>
                {achievements_html}
            </section>
            
            <section class="section contact-section" id="contact">
                <h2 class="section-title" style="color: white;">Get In Touch</h2>
                <p style="font-size: 1.2em; opacity: 0.9; max-width: 700px; margin: 0 auto 20px;">Let's collaborate and create something amazing together!</p>
                <div class="contact-grid">
                    <div class="contact-item"><div class="contact-label">Email</div><div class="contact-value">{data['personal_info'].get('email', '')}</div></div>
                    <div class="contact-item"><div class="contact-label">LinkedIn</div><div class="contact-value">{data['personal_info'].get('linkedin', '').replace('https://', '').replace('http://', '')}</div></div>
                    <div class="contact-item"><div class="contact-label">GitHub</div><div class="contact-value">{data['personal_info'].get('github', '').replace('https://', '').replace('http://', '')}</div></div>
                    {f'<div class="contact-item"><div class="contact-label">Twitter</div><div class="contact-value">{data["personal_info"].get("twitter", "").replace("https://", "").replace("http://", "")}</div></div>' if data['personal_info'].get('twitter') else ''}
                </div>
                <a href="#" class="download-btn" onclick="alert('Generate DOCX from the Download tab!'); return false;">📄 Download Resume</a>
            </section>
        </body></html>
        """,
        
        "Dark Modern Portfolio": "TEMPLATE_2",
        "Neon Cyberpunk": "TEMPLATE_3",
        "Elegant Minimalist": "TEMPLATE_4"
    }
    
    return templates.get(template_name, templates["Designer Portfolio"])

def render_portfolio_builder_page():
    st.title("🎨 Professional Portfolio Builder")
    st.markdown("### Create stunning portfolios with modern, responsive templates")
    
    if 'portfolio_data' not in st.session_state:
        st.session_state.portfolio_data = {
            'personal_info': {
                'name': 'John Doe',
                'tagline': 'Full-Stack Developer & UI/UX Enthusiast',
                'email': 'john@example.com',
                'phone': '+1 (555) 123-4567',
                'linkedin': 'linkedin.com/in/johndoe',
                'github': 'github.com/johndoe',
                'twitter': 'twitter.com/johndoe'
            },
            'about': 'Passionate full-stack developer with 5+ years of experience building scalable web applications. Specialized in React, Node.js, and cloud technologies. Love turning complex problems into elegant solutions.',
            'skills': {
                'Frontend': ['React', 'Vue.js', 'TypeScript', 'Tailwind CSS'],
                'Backend': ['Node.js', 'Python', 'Django', 'PostgreSQL'],
                'Tools': ['Git', 'Docker', 'AWS', 'CI/CD'],
                'Soft Skills': ['Leadership', 'Communication', 'Problem Solving']
            },
            'projects': [
                {
                    'title': 'E-Commerce Platform',
                    'tech_stack': 'React, Node.js, MongoDB, Stripe',
                    'description': 'Built a fully-featured e-commerce platform with payment integration, inventory management, and real-time analytics.',
                    'link': 'https://example.com'
                },
                {
                    'title': 'AI Content Generator',
                    'tech_stack': 'Python, FastAPI, OpenAI, React',
                    'description': 'Developed an AI-powered content generation tool that helps marketers create engaging content.',
                    'link': 'https://example.com'
                }
            ],
            'experience': [
                {
                    'role': 'Senior Full Stack Developer',
                    'company': 'Tech Corp Inc.',
                    'duration': 'Jan 2022 - Present',
                    'description': 'Leading a team of 5 developers in building scalable web applications. Implemented CI/CD pipelines reducing deployment time by 60%.'
                }
            ],
            'education': [
                {
                    'degree': 'B.Tech in Computer Science',
                    'college': 'MIT University',
                    'cgpa': '8.5/10',
                    'duration': '2015 - 2019'
                }
            ],
            'certifications': [
                {
                    'name': 'AWS Certified Solutions Architect',
                    'issuer': 'Amazon Web Services'
                },
                {
                    'name': 'Google Cloud Professional Developer',
                    'issuer': 'Google Cloud'
                }
            ],
            'achievements': [
                'Won Best Innovation Award at TechHack 2023',
                'Published research paper on Machine Learning in IEEE',
                'Mentored 50+ students in web development'
            ]
        }
    
    with st.sidebar:
        st.header("⚙️ Template Settings")
        template_name = st.selectbox(
            "Choose Your Template",
            ["Designer Portfolio", "Dark Modern Portfolio", "Neon Cyberpunk", "Elegant Minimalist"],
            help="Each template has a unique design and layout"
        )
        
        st.markdown("---")
        template_descriptions = {
            "Designer Portfolio": "🎨 Creative layout with vibrant gradients",
            "Dark Modern Portfolio": "🌙 Sleek dark theme with glassmorphism",
            "Neon Cyberpunk": "⚡ Futuristic neon design",
            "Elegant Minimalist": "✨ Clean and sophisticated"
        }
        st.info(template_descriptions[template_name])
        
        st.markdown("---")
        st.markdown("### 📊 Quick Stats")
        st.metric("Projects", len(st.session_state.portfolio_data.get('projects', [])))
        st.metric("Experience", len(st.session_state.portfolio_data.get('experience', [])))
        st.metric("Certifications", len(st.session_state.portfolio_data.get('certifications', [])))
    
    tab1, tab2, tab3 = st.tabs(["✏️ Edit Content", "👁️ Live Preview", "⬇️ Download"])
    
    with tab1:
        with st.expander("📸 Profile Photo", expanded=True):
            col1, col2 = st.columns([3, 1])
            
            with col1:
                uploaded_file = st.file_uploader(
                    "Upload your professional photo",
                    type=['jpg', 'jpeg', 'png'],
                    help="Recommended: Square image, 400x400px minimum"
                )
            
            with col2:
                if st.button("🗑️ Clear", use_container_width=True, type="secondary"):
                    if 'portfolio_photo' in st.session_state:
                        del st.session_state.portfolio_photo
                        st.rerun()
        
        
        with st.expander("👤 Personal Information", expanded=True):
            col1, col2 = st.columns(2)
            with col1:
                name = st.text_input("Full Name *", st.session_state.portfolio_data['personal_info']['name'])
                email = st.text_input("Email *", st.session_state.portfolio_data['personal_info']['email'])
                phone = st.text_input("Phone", st.session_state.portfolio_data['personal_info']['phone'])
            with col2:
                tagline = st.text_input("Professional Tagline *", st.session_state.portfolio_data['personal_info']['tagline'])
                website = st.text_input("Website", st.session_state.portfolio_data['personal_info']['website'])
                github = st.text_input("GitHub", st.session_state.portfolio_data['personal_info']['github'])
        
        with st.expander("📝 About Me", expanded=True):
            about = st.text_area(
                "Tell your story",
                st.session_state.portfolio_data['about'],
                height=120,
                help="Write 2-3 sentences about your background and expertise"
            )
        
        with st.expander("💼 Projects", expanded=True):
            num_projects = st.number_input("Number of Projects", 1, 10, len(st.session_state.portfolio_data['projects']), help="Add 3-5 key projects")
            projects = []
            for i in range(num_projects):
                st.markdown(f"#### 🚀 Project {i+1}")
                col1, col2 = st.columns(2)
                
                default_proj = st.session_state.portfolio_data['projects'][i] if i < len(st.session_state.portfolio_data['projects']) else {}
                
                with col1:
                    title = st.text_input(f"Project Title *", default_proj.get('title', f'Project {i+1}'), key=f"proj_title_{i}")
                    tech = st.text_input(f"Tech Stack", default_proj.get('tech_stack', ''), key=f"proj_tech_{i}", help="e.g., React, Node.js, MongoDB")
                with col2:
                    desc = st.text_area(f"Description", default_proj.get('description', ''), key=f"proj_desc_{i}", height=100)
                    link = st.text_input(f"Project Link", default_proj.get('link', ''), key=f"proj_link_{i}", placeholder="https://...")
                
                projects.append({'title': title, 'tech_stack': tech, 'description': desc, 'link': link})
                if i < num_projects - 1:
                    st.divider()
        
        with st.expander("🎯 Skills & Expertise", expanded=True):
            st.markdown("Add your skills by category")
            col1, col2 = st.columns(2)
            
            default_skills = st.session_state.portfolio_data['skills']
            
            with col1:
                frontend = st.text_area("Frontend Technologies", ', '.join(default_skills.get('Frontend', [])), help="Comma-separated list")
                backend = st.text_area("Backend Technologies", ', '.join(default_skills.get('Backend', [])))
            with col2:
                tools = st.text_area("Tools & Platforms", ', '.join(default_skills.get('Tools', [])))
                soft_skills = st.text_area("Soft Skills", ', '.join(default_skills.get('Soft Skills', [])))
        
        st.session_state.portfolio_data = {
            'personal_info': {
                'name': name, 
                'tagline': tagline, 
                'email': email, 
                'phone': phone, 
                'website': website, 
                'github': github
            },
            'about': about,
            'projects': projects,
            'skills': {
                'Frontend': [s.strip() for s in frontend.split(',') if s.strip()],
                'Backend': [s.strip() for s in backend.split(',') if s.strip()],
                'Tools': [s.strip() for s in tools.split(',') if s.strip()],
                'Soft Skills': [s.strip() for s in soft_skills.split(',') if s.strip()]
            }
        }
        
        if st.button("💾 Save Changes", type="primary", use_container_width=True):
            st.success("✅ Portfolio data saved successfully!")
            st.balloons()
    
    with tab2:
        st.markdown("### 👁️ Live Preview")
        st.info("💡 This is how your portfolio will look. Scroll to see the full design!")
        
        photo_base64 = st.session_state.get('portfolio_photo', None)
        preview_html = generate_preview_html(st.session_state.portfolio_data, template_name, photo_base64)
        st.components.v1.html(preview_html, height=1800, scrolling=True)
    
    with tab3:
        
        st.markdown("### ⬇️ Download Your Portfolio")
                
        # Calculate Portfolio Completeness Score
        data = st.session_state.portfolio_data
        score = 0
        max_score = 100
        feedback = []
        
        # Personal Info (20 points)
        personal_score = 0
        if data['personal_info'].get('name'): personal_score += 5
        if data['personal_info'].get('email'): personal_score += 5
        if data['personal_info'].get('tagline'): personal_score += 5
        if data['personal_info'].get('github'): 
            personal_score += 5
        else:
            feedback.append("❌ Add GitHub profile link")
        score += personal_score
        
        # About Section (10 points)
        if data.get('about') and len(data['about']) > 100:
            score += 10
        else:
            feedback.append("❌ Write a detailed About Me (100+ characters)")
        
        # Projects (30 points)
        project_count = len(data.get('projects', []))
        if project_count >= 3:
            score += 15
        elif project_count >= 2:
            score += 10
            feedback.append("⚠️ Add at least 3 projects (currently: {})".format(project_count))
        else:
            score += 5
            feedback.append("❌ Add at least 3 projects (currently: {})".format(project_count))
        
        # Check project details
        projects_with_tech = sum(1 for p in data.get('projects', []) if p.get('tech_stack'))
        projects_with_links = sum(1 for p in data.get('projects', []) if p.get('link'))
        projects_with_desc = sum(1 for p in data.get('projects', []) if p.get('description') and len(p['description']) > 50)
        
        if projects_with_tech >= project_count * 0.8:
            score += 5
        else:
            feedback.append("⚠️ Add tech stack to all projects ({}/{})".format(projects_with_tech, project_count))
        
        if projects_with_links >= project_count * 0.6:
            score += 5
        else:
            feedback.append("⚠️ Add live demo/GitHub links to projects ({}/{})".format(projects_with_links, project_count))
        
        if projects_with_desc >= project_count * 0.8:
            score += 5
        else:
            feedback.append("⚠️ Write detailed descriptions for all projects")
        
        # Skills (15 points)
        skill_categories = len([v for v in data.get('skills', {}).values() if v])
        total_skills = sum(len(v) for v in data.get('skills', {}).values())
        
        if skill_categories >= 4:
            score += 8
        elif skill_categories >= 3:
            score += 5
            feedback.append("⚠️ Add more skill categories")
        else:
            feedback.append("❌ Add at least 3 skill categories")
        
        if total_skills >= 12:
            score += 7
        elif total_skills >= 8:
            score += 4
            feedback.append("⚠️ Add more skills (currently: {})".format(total_skills))
        else:
            feedback.append("❌ Add at least 12 skills across categories")
        
        # Experience (10 points)
        exp_count = len(data.get('experience', []))
        if exp_count >= 1:
            score += 10
        else:
            feedback.append("❌ Add at least 1 work experience")
        
        # Education (5 points)
        edu_count = len(data.get('education', []))
        if edu_count >= 1:
            score += 5
        else:
            feedback.append("❌ Add your education details")
        
        # Certifications (5 points)
        cert_count = len(data.get('certifications', []))
        if cert_count >= 2:
            score += 5
        elif cert_count >= 1:
            score += 3
            feedback.append("⚠️ Add more certifications")
        else:
            feedback.append("⚠️ Add certifications (optional but recommended)")
        
        # Achievements (5 points)
        achievement_count = len(data.get('achievements', []))
        if achievement_count >= 2:
            score += 5
        elif achievement_count >= 1:
            score += 3
            feedback.append("⚠️ Add more achievements")
        else:
            feedback.append("⚠️ Add notable achievements (optional)")
        
        # Display Score
        score_color = "green" if score >= 80 else "orange" if score >= 60 else "red"
        score_emoji = "🎉" if score >= 80 else "⚡" if score >= 60 else "⚠️"
        
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); 
                    padding: 30px; border-radius: 15px; margin-bottom: 30px; color: white;">
            <h2 style="margin: 0 0 10px 0; color: white;">Portfolio Completeness Score</h2>
            <div style="font-size: 4em; font-weight: 900; margin: 20px 0;">{score_emoji} {score}%</div>
            <div style="background: rgba(255,255,255,0.2); height: 20px; border-radius: 10px; overflow: hidden;">
                <div style="background: {score_color}; height: 100%; width: {score}%; transition: width 0.5s;"></div>
            </div>
            <p style="margin-top: 15px; font-size: 1.1em; opacity: 0.9;">
                {'🌟 Excellent! Your portfolio is ready to impress.' if score >= 80 else 
                 '👍 Good progress! A few improvements will make it outstanding.' if score >= 60 else
                 '💪 Keep building! Follow the suggestions below.'}
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Show Feedback
        if feedback:
            st.markdown("#### 📋 Improvement Suggestions")
            for item in feedback:
                st.markdown(f"- {item}")
            st.markdown("---")
        else:
            st.success("✅ Your portfolio is complete and ready to download!")
        
        
        st.markdown("Choose your preferred format:")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.markdown("#### 📄 DOCX Format")
            st.markdown("*Word document*")
            if st.button("Generate DOCX", type="primary", use_container_width=True):
                with st.spinner("Creating document..."):
                    builder = PortfolioBuilder()
                    buffer = builder.generate_portfolio(st.session_state.portfolio_data, template_name)
                    st.download_button(
                        "⬇️ Download DOCX",
                        data=buffer,
                        file_name=f"portfolio_{template_name.lower().replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
        
        with col2:
            st.markdown("#### 🌐 HTML Format")
            st.markdown("*Ready for web*")
            if st.button("Generate HTML", type="primary", use_container_width=True):
                with st.spinner("Creating HTML..."):
                    photo_base64 = st.session_state.get('portfolio_photo', None)
                    html_content = generate_preview_html(st.session_state.portfolio_data, template_name, photo_base64)
                    st.download_button(
                        "⬇️ Download HTML",
                        data=html_content,
                        file_name=f"portfolio_{template_name.lower().replace(' ', '_')}.html",
                        mime="text/html",
                        use_container_width=True
                    )
        
        with col3:
            st.markdown("#### 📋 JSON Format")
            st.markdown("*Data backup*")
            if st.button("Generate JSON", type="secondary", use_container_width=True):
                import json
                json_data = json.dumps(st.session_state.portfolio_data, indent=2)
                st.download_button(
                    "⬇️ Download JSON",
                    data=json_data,
                    file_name="portfolio_data.json",
                    mime="application/json",
                    use_container_width=True
                )
        
        st.markdown("---")
        st.success("💡 **Pro Tip:** Download HTML for instant web deployment!")
        st.markdown("""
        **What you can do with the files:**
        - 📄 **DOCX**: Edit in Microsoft Word or Google Docs
        - 🌐 **HTML**: Host on GitHub Pages, Netlify, or Vercel
        - 📋 **JSON**: Import your data later or share with others
        """)

if __name__ == "__main__":
    render_portfolio_builder_page()
    
    
    
    
    
    