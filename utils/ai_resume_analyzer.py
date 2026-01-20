import os
from dotenv import load_dotenv
import pdfplumber
import tempfile
import re
from groq import Groq

class AIResumeAnalyzer:
    def __init__(self):
        load_dotenv()
        self.groq_api_key = os.getenv("GROQ_API_KEY")
        
        if not self.groq_api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables. Please add it to your .env file.")
        
        self.client = Groq(api_key=self.groq_api_key)
    
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF using pdfplumber"""
        text = ""
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
            if hasattr(pdf_file, 'getbuffer'):
                temp_file.write(pdf_file.getbuffer())
            elif hasattr(pdf_file, 'read'):
                temp_file.write(pdf_file.read())
                pdf_file.seek(0)
            else:
                temp_file.write(pdf_file)
            temp_path = temp_file.name
        
        try:
            with pdfplumber.open(temp_path) as pdf:
                for page in pdf.pages:
                    try:
                        import warnings
                        with warnings.catch_warnings():
                            warnings.filterwarnings("ignore")
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    except Exception:
                        pass
            
            os.unlink(temp_path)
            return text.strip()
        except Exception:
            try:
                os.unlink(temp_path)
            except:
                pass
            return ""
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        from docx import Document
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(docx_file.getbuffer())
            temp_path = temp_file.name
        
        text = ""
        try:
            doc = Document(temp_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception:
            pass
        
        os.unlink(temp_path)
        return text
    
    def optimize_resume_with_gemini(self, resume_text, job_description=None, job_role=None):
        """Generate ATS optimized resume using GROQ AI"""
        if not resume_text:
            return {"error": "Resume text is required for optimization."}
        
        if not self.groq_api_key:
            return {"error": "GROQ API key is not configured. Please add it to your .env file."}
        
        try:
            prompt = f"""You are an expert resume writer and ATS optimization specialist. Your task is to rewrite the following resume to make it highly ATS-friendly while maintaining all the candidate's information and achievements.

**IMPORTANT FORMATTING INSTRUCTIONS:**
- Structure the resume in clear sections with headers
- Use simple, clean formatting that ATS systems can parse
- Include section headers like: CONTACT INFORMATION, PROFESSIONAL SUMMARY, SKILLS, WORK EXPERIENCE, EDUCATION, CERTIFICATIONS
- Use bullet points (•) for lists
- Keep formatting consistent throughout
- Do NOT use tables, graphics, or complex formatting
- Use standard fonts and simple layouts
- Do NOT use markdown symbols like #, ##, or any other special formatting characters in section headers
- Section headers should be plain uppercase text without any symbols

**OPTIMIZATION REQUIREMENTS:**
1. Add relevant keywords from the job role/description throughout the resume naturally
2. Quantify achievements with specific numbers and metrics where possible
3. Use strong action verbs (e.g., Led, Developed, Implemented, Achieved)
4. Ensure skills section includes both hard and soft skills
5. Make the professional summary compelling and keyword-rich
6. Optimize job titles and descriptions for ATS scanning
7. Include relevant technical skills and tools
8. Ensure dates, locations, and contact information are clearly formatted

**Original Resume:**
{resume_text}
"""
            
            if job_role:
                prompt += f"\n\n**Target Job Role:** {job_role}"
            
            if job_description:
                prompt += f"\n\n**Job Description to optimize for:**\n{job_description}"
            
            prompt += """

**OUTPUT FORMAT:**
Please provide the complete optimized resume in a clean, ATS-friendly format. Make sure to:
- Preserve all the candidate's information
- Enhance the content with relevant keywords
- Use clear section headers WITHOUT any # symbols or markdown formatting
- Format consistently
- Make it visually clean and professional

Return ONLY the optimized resume content without any preamble or explanation."""
            
            # Call GROQ API
            response = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert resume writer and ATS optimization specialist. Provide only the optimized resume content without any preamble, explanation, or markdown symbols like # or ##. Use plain uppercase text for section headers."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model="llama-3.3-70b-versatile",
                temperature=0.7,
                max_tokens=4000,
                top_p=1,
                stream=False
            )
            
            optimized_text = response.choices[0].message.content.strip()
            
            # Remove any # symbols from the output
            optimized_text = re.sub(r'^#+\s*', '', optimized_text, flags=re.MULTILINE)
            
            return {
                "optimized_resume": optimized_text,
                "success": True
            }
        
        except Exception as e:
            return {"error": f"Optimization failed: {str(e)}"}
    
    def analyze_resume_with_gemini(self, resume_text, job_description=None, job_role=None):
        """Analyze resume using GROQ AI"""
        if not resume_text:
            return {"error": "Resume text is required for analysis."}
        
        if not self.groq_api_key:
            return {"error": "GROQ API key is not configured."}
        
        try:
            prompt = f"""You are an expert resume analyst with deep knowledge of industry standards, job requirements, and hiring practices. Analyze the following resume and provide detailed feedback.

Resume:
{resume_text}
"""
            
            if job_role:
                prompt += f"\n\nTarget Role: {job_role}"
            
            if job_description:
                prompt += f"\n\nJob Description: {job_description}"
            
            prompt += """

Provide a comprehensive analysis with the following sections:

Overall Assessment
[Provide a detailed assessment of the resume's overall quality and effectiveness]

Professional Profile Analysis
[Analyze the candidate's professional profile and career trajectory]

Skills Analysis
- **Current Skills**: [List all skills demonstrated]
- **Skill Proficiency**: [Assess expertise levels]
- **Missing Skills**: [List important missing skills]

Experience Analysis
[Analyze how well experience is presented]

Education Analysis
[Analyze education section]

Key Strengths
[List 5-7 specific strengths]

Areas for Improvement
[List 5-7 areas for improvement]

ATS Optimization Assessment
[Analyze ATS compatibility and provide ATS Score: XX/100]

Recommended Courses/Certifications
[Suggest 5-7 relevant courses]

Resume Score
Resume Score: XX/100 [Provide overall score]"""
            
            # Call GROQ API
            response = self.client.chat.completions.create(
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert resume analyst. Provide detailed, structured analysis following the exact format requested."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                model="llama-3.3-70b-versatile",
                temperature=0.7,
                max_tokens=4000,
                top_p=1,
                stream=False
            )
            
            analysis = response.choices[0].message.content.strip()
            
            # Extract scores
            resume_score = self._extract_score(analysis, "Resume Score")
            ats_score = self._extract_score(analysis, "ATS Score")
            
            return {
                "analysis": analysis,
                "resume_score": resume_score,
                "ats_score": ats_score,
                "success": True
            }
        
        except Exception as e:
            return {"error": f"Analysis failed: {str(e)}"}
    
    def _extract_score(self, text, score_type):
        """Extract score from analysis text"""
        try:
            pattern = rf"{score_type}[:\s]+(\d{{1,3}})"
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                score = int(match.group(1))
                return max(0, min(score, 100))  # Ensure score is between 0-100
            return 0
        except:
            return 0
    
    def create_docx_resume(self, optimized_text, candidate_name="Candidate"):
        """Create a professionally formatted DOCX file from optimized resume text"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import io
            import re
            
            # Remove any # symbols from the text before processing
            optimized_text = re.sub(r'^#+\s*', '', optimized_text, flags=re.MULTILINE)
            
            doc = Document()
            
            # Set professional margins (slightly tighter)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.6)
                section.bottom_margin = Inches(0.6)
                section.left_margin = Inches(0.7)
                section.right_margin = Inches(0.7)
            
            # Configure default styles
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Helper function to add horizontal line
            def add_horizontal_line(paragraph, color='4472C4', size='12'):
                p = paragraph._p
                pPr = p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), size)
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), color)
                pBdr.append(bottom)
                pPr.append(pBdr)
            
            # Parse the resume text
            lines = optimized_text.split('\n')
            
            # State tracking
            name_added = False
            contact_section_done = False
            current_section = None
            
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                
                if not line:
                    i += 1
                    continue
                
                # ====================
                # NAME (First non-empty line)
                # ====================
                if not name_added and i < 5 and len(line.split()) <= 5 and not any(char in line for char in ['@', 'http', '|', '•', '(', ')']):
                    name_para = doc.add_paragraph()
                    name_run = name_para.add_run(line.upper())
                    name_run.font.name = 'Calibri'
                    name_run.font.size = Pt(22)
                    name_run.font.bold = True
                    name_run.font.color.rgb = RGBColor(25, 25, 112)  # Midnight Blue
                    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    name_para.paragraph_format.space_after = Pt(4)
                    name_added = True
                    i += 1
                    continue
                
                # ====================
                # CONTACT INFORMATION
                # ====================
                if not contact_section_done and (
                    any(indicator in line.lower() for indicator in ['@', 'phone:', 'email:', 'linkedin.com', 'github.com']) or 
                    re.search(r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}', line) or
                    (i < 8 and '|' in line and not line.isupper())
                ):
                    # Parse contact info
                    contact_para = doc.add_paragraph()
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Split by pipe or common separators
                    if '|' in line:
                        parts = [p.strip() for p in line.split('|')]
                    else:
                        parts = [line.strip()]
                    
                    for idx, part in enumerate(parts):
                        # Create run for each part
                        run = contact_para.add_run(part)
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(70, 70, 70)
                        
                        # Add separator between parts
                        if idx < len(parts) - 1:
                            sep = contact_para.add_run(' • ')
                            sep.font.name = 'Calibri'
                            sep.font.size = Pt(10)
                            sep.font.color.rgb = RGBColor(100, 100, 100)
                    
                    contact_para.paragraph_format.space_after = Pt(2)
                    
                    # Check if this is the last contact line
                    if '@' in line or 'linkedin' in line.lower() or 'github' in line.lower():
                        # Add spacing after contact section
                        space_para = doc.add_paragraph()
                        space_para.paragraph_format.space_after = Pt(10)
                        contact_section_done = True
                    
                    i += 1
                    continue
                
                # ====================
                # SECTION HEADERS
                # ====================
                is_section_header = (
                    line.isupper() and len(line) < 50 and
                    any(keyword in line.upper() for keyword in [
                        'SUMMARY', 'OBJECTIVE', 'EXPERIENCE', 'EDUCATION', 
                        'SKILLS', 'CERTIFICATIONS', 'CERTIFICATION', 'PROJECTS', 
                        'ACHIEVEMENTS', 'PROFESSIONAL', 'WORK', 'TECHNICAL',
                        'PROJECT EXPERIENCE'
                    ]) and 
                    not line.startswith('•') and 
                    not line.startswith('-') and 
                    not line.startswith('*') and
                    'CONTACT INFORMATION' not in line.upper()
                )
                
                if is_section_header:
                    # Add some space before new section
                    if current_section is not None:
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(4)
                    
                    # Section header
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(line.upper())
                    header_run.font.name = 'Calibri'
                    header_run.font.size = Pt(13)
                    header_run.font.bold = True
                    header_run.font.color.rgb = RGBColor(25, 25, 112)  # Midnight Blue
                    
                    # Add thicker horizontal line
                    add_horizontal_line(header_para, color='4472C4', size='16')
                    
                    header_para.paragraph_format.space_before = Pt(2)
                    header_para.paragraph_format.space_after = Pt(6)
                    header_para.paragraph_format.keep_with_next = True
                    
                    current_section = line.upper()
                    i += 1
                    continue
                
                # ====================
                # BULLET POINTS
                # ====================
                if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                    clean_line = line.lstrip('•-* ').strip()
                    
                    bullet_para = doc.add_paragraph(style='List Bullet')
                    bullet_run = bullet_para.add_run(clean_line)
                    bullet_run.font.name = 'Calibri'
                    bullet_run.font.size = Pt(10.5)
                    bullet_run.font.color.rgb = RGBColor(40, 40, 40)
                    
                    bullet_format = bullet_para.paragraph_format
                    bullet_format.space_before = Pt(1)
                    bullet_format.space_after = Pt(3)
                    bullet_format.line_spacing = 1.15
                    bullet_format.left_indent = Inches(0.25)
                    bullet_format.first_line_indent = Inches(-0.2)
                    
                    i += 1
                    continue
                
                # ====================
                # JOB TITLE / COMPANY LINES
                # ====================
                has_date = re.search(r'\b\d{4}\b|\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\b', line)
                looks_like_title = len(line) < 120 and (has_date or '|' in line or '–' in line or ' - ' in line)
                
                if looks_like_title and not line.lower().startswith('note'):
                    para = doc.add_paragraph()
                    
                    # Try to parse: Title | Company | Location | Date
                    # or: Title - Company - Date
                    parts = re.split(r'\s*[\|–]\s*|\s+-\s+', line)
                    
                    for idx, part in enumerate(parts):
                        part = part.strip()
                        if not part:
                            continue
                        
                        run = para.add_run(part)
                        run.font.name = 'Calibri'
                        
                        # First part = Job Title (bold)
                        if idx == 0:
                            run.font.size = Pt(11.5)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        # Last part with date (italic)
                        elif re.search(r'\b\d{4}\b', part):
                            run.font.size = Pt(10)
                            run.font.italic = True
                            run.font.color.rgb = RGBColor(90, 90, 90)
                        # Company/Location (regular)
                        else:
                            run.font.size = Pt(10.5)
                            run.font.color.rgb = RGBColor(60, 60, 60)
                        
                        # Add separator
                        if idx < len(parts) - 1:
                            sep_run = para.add_run(' | ')
                            sep_run.font.name = 'Calibri'
                            sep_run.font.size = Pt(10)
                            sep_run.font.color.rgb = RGBColor(130, 130, 130)
                    
                    para.paragraph_format.space_before = Pt(8)
                    para.paragraph_format.space_after = Pt(3)
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.keep_with_next = True
                    
                    i += 1
                    continue
                
                # ====================
                # REGULAR PARAGRAPHS
                # ====================
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
                
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(4)
                para.paragraph_format.line_spacing = 1.15
                
                i += 1
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
        
        except Exception as e:
            import streamlit as st
            st.error(f"Error creating DOCX: {str(e)}")
            import traceback
            st.code(traceback.format_exc())
            return None